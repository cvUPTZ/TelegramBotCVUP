from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_or_get_style(styles, name, style_type):
    try:
        return styles.add_style(name, style_type)
    except ValueError:
        return styles[name]

def create_docx_memo():
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = create_or_get_style(styles, 'CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.color.rgb = RGBColor(0, 51, 102)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = create_or_get_style(styles, 'CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(18)
    h1_font.color.rgb = RGBColor(0, 51, 102)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = create_or_get_style(styles, 'CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(16)
    h2_font.color.rgb = RGBColor(0, 102, 204)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Normal text style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Add content to the document
    doc.add_paragraph('إنشاء نماذج مذكرات تربوية', style='CustomTitle')
    
    doc.add_paragraph('معلومات عامة', style='CustomHeading1')
    info = doc.add_paragraph(style='Normal')
    info.add_run('المادة: ').bold = True
    info.add_run('تاريخ')
    info.add_run('\nالمستوى: ').bold = True
    info.add_run('السنة الثانية جميع الشعب')
    info.add_run('\nالمدة: ').bold = True
    info.add_run('3 ساعات كل الشعب')
    info.add_run('\nطريقة العمل: ').bold = True
    info.add_run('فردية و جماعية')
    
    doc.add_paragraph('الوحدة التعلمية الأولى', style='CustomHeading1')
    doc.add_paragraph('الاستعمار الأوربي في آسيا و إفريقيا و مقاومته 1815 ـ 1954', style='Normal')
    
    doc.add_paragraph('الكفاءة القاعدية 01', style='CustomHeading2')
    doc.add_paragraph('أمام وضعيات إشكالية تعكس ظاهرة الحركة الاستعمارية الأوربية و حيثيات ظهورها و توسعها في قارتي آسيا و إفريقيا و ما قابلها من كفاح تحرري يكون المتعلم قادرا على:', style='Normal')
    
    abilities = doc.add_paragraph(style='Normal')
    abilities.add_run('• تصنيف\n• شرح\n• تحليل (الظروف، الأسباب، الأهداف، الوسائل، الانعكاسات)')
    
    doc.add_paragraph('اعتمادا على مختلف السندات التاريخية ذات الدلالة.', style='Normal')
    
    doc.add_paragraph('الوضعية التعلمية الثانية', style='CustomHeading1')
    doc.add_paragraph('الحركة الاستعمارية (أسباب ـ أهداف)', style='Normal')
    
    doc.add_paragraph('الكفاءة المستهدفة', style='CustomHeading2')
    doc.add_paragraph('يتعرف على الحركة الاستعمارية الأوروبية و يدرك انعكاساتها و مخاطرها.', style='Normal')
    
    doc.add_paragraph('الإشكالية', style='CustomHeading2')
    doc.add_paragraph('وظفت أوروبا تفوقها الاقتصادي و العسكري لفرض سيطرتها الاستعمارية على شعوب إفريقيا و آسيا. فما هي أسباب و أهداف هذه السيطرة؟', style='Normal')
    
    doc.add_paragraph('خطة الدرس', style='CustomHeading1')
    
    doc.add_paragraph('1. الظروف التاريخية لنشأة الحركة الاستعمارية (1 ساعة 30 دقيقة)', style='CustomHeading2')
    circumstances = doc.add_paragraph(style='Normal')
    circumstances.add_run('• النهضة الأوروبية و تطور أوروبا\n• الثورة الصناعية و تفوق أوروبا اقتصاديا و عسكريا\n• الكشوفات الجغرافية و نتائجها\n• التنافس بين الدول الأوروبية\n• الصراع الحضاري بين العالم الإسلامي و الغرب المسيحي\n• ضعف شعوب إفريقيا و آسيا')
    
    doc.add_paragraph('2. أسباب الحركة الاستعمارية و أهدافها (1 ساعة 30 دقيقة)', style='CustomHeading2')
    doc.add_paragraph('أ. الأسباب:', style='Normal')
    reasons = doc.add_paragraph(style='Normal')
    reasons.add_run('• اقتصادية\n• بشرية\n• إستراتيجية\n• سياسية\n• حضارية')
    
    doc.add_paragraph('ب. الأهداف:', style='Normal')
    goals = doc.add_paragraph(style='Normal')
    goals.add_run('• اقتصادية و مالية\n• إستراتيجية و عسكرية\n• حضارية و دينية')
    
    doc.add_paragraph('3. تغلغل الاستعمار الأوروبي في إفريقيا و آسيا', style='CustomHeading2')
    doc.add_paragraph('الطرق:', style='Normal')
    methods = doc.add_paragraph(style='Normal')
    methods.add_run('• مباشرة\n• غير مباشرة')
    
    doc.add_paragraph('وسائل تنفيذ المخططات الاستعمارية:', style='Normal')
    means = doc.add_paragraph(style='Normal')
    means.add_run('• الامتيازات\n• دور البنوك و الشركات\n• الاتفاقيات الودية\n• الاتفاقيات و الوعود السرية')
    
    doc.add_paragraph('تقويم مرحلي', style='CustomHeading1')
    doc.add_paragraph('اكتب فقرة تتناول من خاللها خطورة الفكر الاستعماري الأوروبي على شعوب المستعمرات.', style='Normal')
    
    # Add a border to the page
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Save the document
    doc.save('educational_memo.docx')
    print("DOCX file 'educational_memo.docx' has been created successfully.")

# Generate the DOCX memo
create_docx_memo()